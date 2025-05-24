"""File handling utilities."""

import json
import yaml
import re
from pathlib import Path
from typing import Any, Dict, Union, List, Tuple
from docx import Document as DocxDocument
import markdown


class FileHandler:
    """Handles reading and writing various file formats."""
    
    def read_file(self, file_path: Union[str, Path]) -> str:
        """Read text content from various file formats."""
        path = Path(file_path)
        
        if path.suffix.lower() == '.docx':
            return self._read_docx(path)
        elif path.suffix.lower() == '.md':
            return self._read_markdown(path)
        else:
            # Default to plain text
            return path.read_text(encoding='utf-8')
    
    def write_file(self, file_path: Union[str, Path], content: str) -> None:
        """Write content to file."""
        path = Path(file_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content, encoding='utf-8')
    
    def read_json(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Read JSON file."""
        path = Path(file_path)
        with path.open('r', encoding='utf-8') as f:
            return json.load(f)
    
    def write_json(self, file_path: Union[str, Path], data: Dict[str, Any]) -> None:
        """Write JSON file."""
        path = Path(file_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        with path.open('w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    
    def read_yaml(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Read YAML file."""
        path = Path(file_path)
        with path.open('r', encoding='utf-8') as f:
            return yaml.safe_load(f)
    
    def write_yaml(self, file_path: Union[str, Path], data: Dict[str, Any]) -> None:
        """Write YAML file."""
        path = Path(file_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        with path.open('w', encoding='utf-8') as f:
            yaml.dump(data, f, default_flow_style=False)
    
    def detect_chapters(self, content: str) -> List[Tuple[str, str]]:
        """
        Detect chapters in text content and return list of (title, content) tuples.
        Handles various chapter title formats and normalizes them.
        """
        from ..core.document import normalize_chapter_title
        
        # Chapter detection patterns (in order of preference)
        chapter_patterns = [
            # "Chapter One", "Chapter 1", etc.
            r'\n\s*(Chapter\s+\w+(?:-\w+)?)\s*[:\.]?\s*\n',
            # "Ch. One", "Ch 1", etc.
            r'\n\s*(Ch\.?\s+\w+(?:-\w+)?)\s*[:\.]?\s*\n',
            # Just numbers: "1.", "2.", etc.
            r'\n\s*(\d+)\.\s*\n',
            # Standalone chapter words: "One", "Two", etc.
            r'\n\s*(One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten|Eleven|Twelve|Thirteen|Fourteen|Fifteen|Sixteen|Seventeen|Eighteen|Nineteen|Twenty)\s*\n',
        ]
        
        best_splits = None
        best_pattern = None
        
        # Try each pattern and use the one that gives the most reasonable splits
        for pattern in chapter_patterns:
            splits = re.split(pattern, content, flags=re.MULTILINE | re.IGNORECASE)
            
            # Skip if we don't get multiple parts
            if len(splits) < 3:
                continue
                
            # Check if this looks like a good chapter split
            # (should have alternating titles and content)
            if len(splits) % 2 == 1:  # Odd number means pattern worked
                best_splits = splits
                best_pattern = pattern
                break
        
        if not best_splits:
            # No chapter patterns found, return the whole content as one chapter
            return [("Chapter 1", content.strip())]
        
        chapters = []
        
        # Process the splits
        # splits[0] is content before first chapter (usually empty or intro)
        # splits[1] is first chapter title, splits[2] is first chapter content
        # splits[3] is second chapter title, splits[4] is second chapter content, etc.
        
        for i in range(1, len(best_splits), 2):
            if i + 1 < len(best_splits):
                raw_title = best_splits[i].strip()
                chapter_content = best_splits[i + 1].strip()
                
                # Normalize the title
                if raw_title.isdigit():
                    # Handle standalone numbers
                    normalized_title = f"Chapter {raw_title}"
                else:
                    normalized_title = normalize_chapter_title(raw_title)
                
                if chapter_content:  # Only add if there's actual content
                    chapters.append((normalized_title, chapter_content))
        
        # If no chapters were found, return the whole content
        if not chapters:
            return [("Chapter 1", content.strip())]
            
        return chapters
    
    def import_chapters_from_file(self, file_path: Union[str, Path]) -> List[Tuple[str, str]]:
        """
        Import chapters from a file, detecting and normalizing chapter titles.
        Returns list of (title, content) tuples.
        """
        content = self.read_file(file_path)
        return self.detect_chapters(content)
    
    def _read_docx(self, path: Path) -> str:
        """Read DOCX file."""
        doc = DocxDocument(path)
        return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
    
    def _read_markdown(self, path: Path) -> str:
        """Read Markdown file."""
        return path.read_text(encoding='utf-8')
    
    def export_chapters_to_file(
        self, 
        chapters: Dict[int, Any], 
        file_path: Union[str, Path], 
        format_type: str = "markdown"
    ) -> None:
        """
        Export chapters to a file in the specified format.
        
        Args:
            chapters: Dictionary of chapter number to chapter object
            file_path: Path to output file
            format_type: Format to export ("markdown", "text", "docx")
        """
        path = Path(file_path)
        
        if format_type.lower() == "markdown":
            self._export_as_markdown(chapters, path)
        elif format_type.lower() == "text":
            self._export_as_text(chapters, path)
        elif format_type.lower() == "docx":
            self._export_as_docx(chapters, path)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
    
    def _export_as_markdown(self, chapters: Dict[int, Any], path: Path) -> None:
        """Export chapters as Markdown file."""
        content = []
        
        for chapter_num in sorted(chapters.keys()):
            chapter = chapters[chapter_num]
            content.append(f"# {chapter.title}\n")
            content.append(f"{chapter.content}\n\n")
            
            # Add scenes if any
            if hasattr(chapter, 'scenes') and chapter.scenes:
                for scene_num in sorted(chapter.scenes.keys()):
                    scene = chapter.scenes[scene_num]
                    if scene.content:
                        content.append(f"## Scene {scene_num}: {scene.title}\n")
                        content.append(f"{scene.content}\n\n")
        
        self.write_file(path, ''.join(content))
    
    def _export_as_text(self, chapters: Dict[int, Any], path: Path) -> None:
        """Export chapters as plain text file."""
        content = []
        
        for chapter_num in sorted(chapters.keys()):
            chapter = chapters[chapter_num]
            content.append(f"{chapter.title}\n")
            content.append("=" * len(chapter.title) + "\n\n")
            content.append(f"{chapter.content}\n\n")
            
            # Add scenes if any
            if hasattr(chapter, 'scenes') and chapter.scenes:
                for scene_num in sorted(chapter.scenes.keys()):
                    scene = chapter.scenes[scene_num]
                    if scene.content:
                        content.append(f"Scene {scene_num}: {scene.title}\n")
                        content.append("-" * (len(scene.title) + 10) + "\n")
                        content.append(f"{scene.content}\n\n")
        
        self.write_file(path, ''.join(content))
    
    def _export_as_docx(self, chapters: Dict[int, Any], path: Path) -> None:
        """Export chapters as DOCX file."""
        doc = DocxDocument()
        
        for chapter_num in sorted(chapters.keys()):
            chapter = chapters[chapter_num]
            
            # Add chapter title as heading
            doc.add_heading(chapter.title, level=1)
            
            # Add chapter content
            if chapter.content:
                doc.add_paragraph(chapter.content)
            
            # Add scenes if any
            if hasattr(chapter, 'scenes') and chapter.scenes:
                for scene_num in sorted(chapter.scenes.keys()):
                    scene = chapter.scenes[scene_num]
                    if scene.content:
                        doc.add_heading(f"Scene {scene_num}: {scene.title}", level=2)
                        doc.add_paragraph(scene.content)
            
            # Add page break after each chapter (except the last one)
            if chapter_num != max(chapters.keys()):
                doc.add_page_break()
        
        doc.save(path)